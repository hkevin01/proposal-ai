"""
Proposal Editor GUI Integration
Integrates AI proposal generation into the main PyQt interface
"""
import json
import sys
import threading
import time

from PyQt5.QtCore import Qt, QThread, pyqtSignal
from PyQt5.QtGui import QFont, QTextDocument
from PyQt5.QtWidgets import (
    QApplication,
    QComboBox,
    QFileDialog,
    QFormLayout,
    QGroupBox,
    QHBoxLayout,
    QLabel,
    QListWidget,
    QMessageBox,
    QProgressBar,
    QPushButton,
    QSpinBox,
    QSplitter,
    QTabWidget,
    QTextEdit,
    QVBoxLayout,
    QWidget,
)

from ..core.database import DatabaseManager
from ..proposals.ai_proposal_generator import (
    AIProposalGenerator,
    ProposalContext,
    ProposalManager,
)


class ProposalGenerationWorker(QThread):
    """Worker thread for AI proposal generation"""
    progress = pyqtSignal(str)
    section_completed = pyqtSignal(str, str)  # section_name, content
    finished = pyqtSignal(dict)  # complete proposal
    error = pyqtSignal(str)
    
    def __init__(self, context, template=None, ai_generator=None):
        super().__init__()
        self.context = context
        self.template = template
        self.ai_generator = ai_generator or AIProposalGenerator()
    
    def run(self):
        try:
            self.progress.emit("Analyzing opportunity requirements...")
            
            # Get or suggest template
            if not self.template:
                self.template = self.ai_generator.suggest_template(self.context)
            # Generate sections one by one
            proposal_sections = {}
            total_sections = len(self.template.sections)
            
            for i, section in enumerate(self.template.sections):
                self.progress.emit(f"Generating {section}... ({i+1}/{total_sections})")
                
                content = self.ai_generator._generate_section_content(
                    section, self.context, self.template
                )
                
                proposal_sections[section] = content
                self.section_completed.emit(section, content)
            
            # Create complete proposal
            proposal = {
                "title": self.context.opportunity_title,
                "organization": self.context.organization,
                "template_used": self.template.name,
                "sections": proposal_sections,
                "word_count": sum(len(content.split()) for content in proposal_sections.values()),
                "word_limit": self.template.word_limit,
                "requirements_met": self.ai_generator._check_requirements(proposal_sections, self.template),
                "suggestions": self.ai_generator._generate_improvement_suggestions(
                    proposal_sections, self.context, self.template
                )
            }
            
            self.finished.emit(proposal)
            
        except Exception as e:
            self.error.emit(str(e))


class ProposalEditorWidget(QWidget):
    """Main proposal editor widget"""
    
    def __init__(self, db_manager=None):
        super().__init__()
        self.db_manager = db_manager or DatabaseManager()
        self.ai_generator = AIProposalGenerator()
        self.proposal_manager = ProposalManager(self.db_manager, self.ai_generator)
        self.current_proposal = None
        self.generation_worker = None
        
        self.setup_ui()
    
    def setup_ui(self):
        """Setup the proposal editor interface"""
        layout = QVBoxLayout(self)
        
        # Control panel
        control_panel = self.create_control_panel()
        layout.addWidget(control_panel)
        
        # Main editor area
        main_splitter = QSplitter(Qt.Horizontal)
        
        # Left panel: Sections and templates
        left_panel = self.create_left_panel()
        main_splitter.addWidget(left_panel)
        
        # Right panel: Editor and preview
        right_panel = self.create_right_panel()
        main_splitter.addWidget(right_panel)
        
        main_splitter.setSizes([300, 700])
        layout.addWidget(main_splitter)
        
        # Status panel
        status_panel = self.create_status_panel()
        layout.addWidget(status_panel)
    
    def create_control_panel(self):
        """Create the top control panel"""
        group = QGroupBox("Proposal Generation")
        layout = QHBoxLayout(group)
        
        # Opportunity selection
        layout.addWidget(QLabel("Opportunity:"))
        self.opportunity_combo = QComboBox()
        self.load_opportunities()
        layout.addWidget(self.opportunity_combo)
        
        # Template selection
        layout.addWidget(QLabel("Template:"))
        self.template_combo = QComboBox()
        self.load_templates()
        layout.addWidget(self.template_combo)
        
        # Generation controls
        self.generate_btn = QPushButton("🤖 Generate Proposal")
        self.generate_btn.clicked.connect(self.generate_proposal)
        layout.addWidget(self.generate_btn)
        
        self.improve_btn = QPushButton("✨ Improve Section")
        self.improve_btn.clicked.connect(self.improve_current_section)
        self.improve_btn.setEnabled(False)
        layout.addWidget(self.improve_btn)
        
        layout.addStretch()
        
        # Export controls
        self.export_pdf_btn = QPushButton("📄 Export PDF")
        self.export_pdf_btn.clicked.connect(self.export_pdf)
        self.export_pdf_btn.setEnabled(False)
        layout.addWidget(self.export_pdf_btn)
        
        self.export_word_btn = QPushButton("📝 Export Word")
        self.export_word_btn.clicked.connect(self.export_word)
        self.export_word_btn.setEnabled(False)
        layout.addWidget(self.export_word_btn)
        
        return group
    
    def create_left_panel(self):
        """Create the left panel with sections and info"""
        widget = QWidget()
        layout = QVBoxLayout(widget)
        
        # Sections list
        sections_group = QGroupBox("Proposal Sections")
        sections_layout = QVBoxLayout(sections_group)
        
        self.sections_list = QListWidget()
        self.sections_list.itemClicked.connect(self.load_section)
        sections_layout.addWidget(self.sections_list)
        
        layout.addWidget(sections_group)
        
        # Requirements checklist
        req_group = QGroupBox("Requirements Status")
        req_layout = QVBoxLayout(req_group)
        
        self.requirements_list = QListWidget()
        req_layout.addWidget(self.requirements_list)
        
        layout.addWidget(req_group)
        
        # Proposal stats
        stats_group = QGroupBox("Statistics")
        stats_layout = QFormLayout(stats_group)
        
        self.word_count_label = QLabel("0")
        self.word_limit_label = QLabel("N/A")
        self.completion_label = QLabel("0%")
        
        stats_layout.addRow("Word Count:", self.word_count_label)
        stats_layout.addRow("Word Limit:", self.word_limit_label)
        stats_layout.addRow("Completion:", self.completion_label)
        
        layout.addWidget(stats_group)
        
        return widget
    
    def create_right_panel(self):
        """Create the right panel with editor and preview"""
        tab_widget = QTabWidget()
        
        # Editor tab
        editor_widget = QWidget()
        editor_layout = QVBoxLayout(editor_widget)
        
        # Section title
        self.section_title_label = QLabel("Select a section to edit")
        self.section_title_label.setFont(QFont("Arial", 14, QFont.Bold))
        editor_layout.addWidget(self.section_title_label)
        
        # Text editor
        self.text_editor = QTextEdit()
        self.text_editor.textChanged.connect(self.update_word_count)
        editor_layout.addWidget(self.text_editor)
        
        # Section controls
        section_controls = QHBoxLayout()
        
        self.regenerate_section_btn = QPushButton("🔄 Regenerate Section")
        self.regenerate_section_btn.clicked.connect(self.regenerate_current_section)
        section_controls.addWidget(self.regenerate_section_btn)
        
        section_controls.addStretch()
        
        editor_layout.addLayout(section_controls)
        
        tab_widget.addTab(editor_widget, "📝 Editor")
        
        # Preview tab
        preview_widget = QWidget()
        preview_layout = QVBoxLayout(preview_widget)
        
        self.preview_text = QTextEdit()
        self.preview_text.setReadOnly(True)
        preview_layout.addWidget(self.preview_text)
        
        refresh_preview_btn = QPushButton("🔄 Refresh Preview")
        refresh_preview_btn.clicked.connect(self.update_preview)
        preview_layout.addWidget(refresh_preview_btn)
        
        tab_widget.addTab(preview_widget, "👁 Preview")
        
        return tab_widget
    
    def create_status_panel(self):
        """Create the bottom status panel"""
        group = QGroupBox("Generation Status")
        layout = QVBoxLayout(group)
        
        self.status_label = QLabel("Ready to generate proposal")
        layout.addWidget(self.status_label)
        
        self.progress_bar = QProgressBar()
        self.progress_bar.setVisible(False)
        layout.addWidget(self.progress_bar)
        
        return group
    
    def load_opportunities(self):
        """Load opportunities from database"""
        self.opportunity_combo.clear()
        self.opportunity_combo.addItem("Select an opportunity...", None)
        
        events = self.db_manager.get_events()
        for event in events:
            # event structure: (id, name, org_id, date, deadline, desc, url, req, status, created, org_name)
            display_text = f"{event[1]} - {event[10] if len(event) > 10 else 'Unknown Org'}"
            self.opportunity_combo.addItem(display_text, event)
    
    def load_templates(self):
        """Load proposal templates"""
        self.template_combo.clear()
        self.template_combo.addItem("Auto-select template", None)
        
        templates = self.proposal_manager.get_proposal_templates()
        for template in templates:
            self.template_combo.addItem(template.name, template)
    
    def generate_proposal(self):
        """Generate a new proposal"""
        # Get selected opportunity
        opportunity_data = self.opportunity_combo.currentData()
        if not opportunity_data:
            QMessageBox.warning(self, "Warning", "Please select an opportunity first.")
            return
        
        # Create context
        context = ProposalContext(
            opportunity_title=opportunity_data[1],
            organization=opportunity_data[10] if len(opportunity_data) > 10 else "Unknown",
            deadline=opportunity_data[4] or "Not specified",
            requirements=opportunity_data[7] or "No specific requirements",
            description=opportunity_data[5] or "No description available",
            keywords=[]  # Could extract from other fields
        )
        
        # Get selected template
        template = self.template_combo.currentData()
        
        # Start generation in worker thread
        self.start_generation(context, template)
    
    def start_generation(self, context, template=None):
        """Start AI generation in worker thread"""
        if self.generation_worker and self.generation_worker.isRunning():
            return
        
        # Update UI
        self.generate_btn.setEnabled(False)
        self.progress_bar.setVisible(True)
        self.progress_bar.setRange(0, 0)  # Indeterminate
        
        # Start worker
        self.generation_worker = ProposalGenerationWorker(context, template, self.ai_generator)
        self.generation_worker.progress.connect(self.update_generation_status)
        self.generation_worker.section_completed.connect(self.add_section_content)
        self.generation_worker.finished.connect(self.generation_completed)
        self.generation_worker.error.connect(self.generation_error)
        self.generation_worker.start()
    
    def update_generation_status(self, message):
        """Update generation status"""
        self.status_label.setText(message)
    
    def add_section_content(self, section_name, content):
        """Add completed section content"""
        # Add to sections list if not already there
        for i in range(self.sections_list.count()):
            if self.sections_list.item(i).text() == section_name:
                self.sections_list.item(i).setData(Qt.UserRole, content)
                return
        
        # Add new section
        item = self.sections_list.addItem(section_name)
        self.sections_list.item(self.sections_list.count() - 1).setData(Qt.UserRole, content)
    
    def generation_completed(self, proposal):
        """Handle completed proposal generation"""
        self.current_proposal = proposal
        
        # Update UI
        self.generate_btn.setEnabled(True)
        self.export_pdf_btn.setEnabled(True)
        self.export_word_btn.setEnabled(True)
        self.improve_btn.setEnabled(True)
        self.progress_bar.setVisible(False)
        
        # Update sections list
        self.sections_list.clear()
        for section_name, content in proposal["sections"].items():
            item = self.sections_list.addItem(section_name)
            self.sections_list.item(self.sections_list.count() - 1).setData(Qt.UserRole, content)
        
        # Update requirements
        self.update_requirements_status(proposal)
        
        # Update stats
        self.update_statistics(proposal)
        
        # Update preview
        self.update_preview()
        
        self.status_label.setText("Proposal generation completed!")
        
        QMessageBox.information(self, "Success", 
                              f"Proposal generated successfully!\n"
                              f"Word count: {proposal['word_count']}\n"
                              f"Sections: {len(proposal['sections'])}")
    
    def generation_error(self, error_message):
        """Handle generation errors"""
        self.generate_btn.setEnabled(True)
        self.progress_bar.setVisible(False)
        self.status_label.setText(f"Error: {error_message}")
        
        QMessageBox.critical(self, "Generation Error", 
                           f"Failed to generate proposal:\n{error_message}")
    
    def load_section(self, item):
        """Load selected section into editor"""
        section_name = item.text()
        content = item.data(Qt.UserRole) or ""
        
        self.section_title_label.setText(f"Editing: {section_name}")
        self.text_editor.setPlainText(content)
    
    def update_word_count(self):
        """Update word count display"""
        text = self.text_editor.toPlainText()
        word_count = len(text.split()) if text.strip() else 0
        
        # Update the current section's word count
        current_item = self.sections_list.currentItem()
        if current_item:
            current_item.setData(Qt.UserRole, text)
        
        # Update total statistics if we have a proposal
        if self.current_proposal:
            total_words = 0
            for i in range(self.sections_list.count()):
                item = self.sections_list.item(i)
                section_text = item.data(Qt.UserRole) or ""
                total_words += len(section_text.split()) if section_text.strip() else 0
            
            self.word_count_label.setText(str(total_words))
            
            # Update completion percentage
            if self.current_proposal.get("word_limit"):
                completion = min(100, (total_words / self.current_proposal["word_limit"]) * 100)
                self.completion_label.setText(f"{completion:.1f}%")
    
    def update_requirements_status(self, proposal):
        """Update requirements checklist"""
        self.requirements_list.clear()
        
        if "requirements_met" in proposal:
            for req in proposal["requirements_met"]:
                self.requirements_list.addItem(f"✅ {req}")
        
        if "suggestions" in proposal:
            for suggestion in proposal["suggestions"]:
                self.requirements_list.addItem(f"💡 {suggestion}")
    
    def update_statistics(self, proposal):
        """Update proposal statistics"""
        self.word_count_label.setText(str(proposal.get("word_count", 0)))
        self.word_limit_label.setText(str(proposal.get("word_limit", "N/A")))
        
        if proposal.get("word_limit"):
            completion = min(100, (proposal["word_count"] / proposal["word_limit"]) * 100)
            self.completion_label.setText(f"{completion:.1f}%")
    
    def update_preview(self):
        """Update preview with full proposal"""
        if not self.current_proposal:
            self.preview_text.setPlainText("No proposal generated yet.")
            return
        
        # Generate formatted preview
        preview_text = f"""
{self.current_proposal['title']}
{'=' * len(self.current_proposal['title'])}

Organization: {self.current_proposal['organization']}
Template: {self.current_proposal['template_used']}
Word Count: {self.current_proposal['word_count']}

"""
        
        # Add all sections
        for i in range(self.sections_list.count()):
            item = self.sections_list.item(i)
            section_name = item.text()
            content = item.data(Qt.UserRole) or ""
            
            preview_text += f"\n{section_name}\n"
            preview_text += "-" * len(section_name) + "\n"
            preview_text += content + "\n\n"
        
        self.preview_text.setPlainText(preview_text)
    
    def improve_current_section(self):
        """Improve the currently selected section using AI"""
        current_item = self.sections_list.currentItem()
        if not current_item:
            QMessageBox.warning(self, "Warning", "Please select a section to improve.")
            return
        section_name = current_item.text()
        current_content = current_item.data(Qt.UserRole) or ""
        if not current_content.strip():
            QMessageBox.warning(self, "Warning", "Section is empty. Nothing to improve.")
            return
        # Use AI to improve the section
        improved = self.ai_generator.improve_section_content(
            section_name, current_content, self.current_proposal, self.template_combo.currentData()
        )
        current_item.setData(Qt.UserRole, improved)
        self.text_editor.setPlainText(improved)
        self.update_preview()
        QMessageBox.information(self, "Improved", "Section improved using AI.")

    def regenerate_current_section(self):
        """Regenerate the currently selected section using AI"""
        current_item = self.sections_list.currentItem()
        if not current_item:
            QMessageBox.warning(self, "Warning", "Please select a section to regenerate.")
            return
        section_name = current_item.text()
        regenerated = self.ai_generator._generate_section_content(
            section_name, self.current_proposal, self.template_combo.currentData()
        )
        current_item.setData(Qt.UserRole, regenerated)
        self.text_editor.setPlainText(regenerated)
        self.update_preview()
        QMessageBox.information(self, "Regenerated", "Section regenerated using AI.")

    def export_pdf(self):
        """Export proposal to PDF"""
        if not self.current_proposal:
            QMessageBox.warning(self, "Warning", "No proposal to export.")
            return
        filename, _ = QFileDialog.getSaveFileName(
            self, "Export PDF", f"{self.current_proposal['title']}.pdf", "PDF Files (*.pdf)"
        )
        if filename:
            from fpdf import FPDF
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=12)
            pdf.multi_cell(0, 10, self.preview_text.toPlainText())
            pdf.output(filename)
            QMessageBox.information(self, "Exported", f"PDF exported to: {filename}")

    def export_word(self):
        """Export proposal to Word document"""
        if not self.current_proposal:
            QMessageBox.warning(self, "Warning", "No proposal to export.")
            return
        filename, _ = QFileDialog.getSaveFileName(
            self, "Export Word", f"{self.current_proposal['title']}.docx", "Word Documents (*.docx)"
        )
        if filename:
            from docx import Document
            doc = Document()
            doc.add_heading(self.current_proposal['title'], 0)
            doc.add_paragraph(f"Organization: {self.current_proposal['organization']}")
            doc.add_paragraph(f"Template: {self.current_proposal['template_used']}")
            doc.add_paragraph(f"Word Count: {self.current_proposal['word_count']}")
            for i in range(self.sections_list.count()):
                item = self.sections_list.item(i)
                doc.add_heading(item.text(), level=1)
                doc.add_paragraph(item.data(Qt.UserRole) or "")
            doc.save(filename)
            QMessageBox.information(self, "Exported", f"Word document exported to: {filename}")


class CollaborativeProposalEditor(ProposalEditorWidget):
    """
    Collaborative Proposal Editor supporting multi-user real-time editing.
    Future: Integrate with web/mobile clients via API.
    """
    def __init__(self, db_manager=None, user_id=None):
        super().__init__(db_manager)
        self.user_id = user_id
        self.active_users = []  # Track active editors
        self.version_history = []  # Store proposal versions
        self.setup_collaboration()

    def setup_collaboration(self):
        """Initialize collaboration features."""
        # Simulate websocket/API setup for real-time sync
        self.collab_active = True
        self.roles = {"editor": [], "viewer": []}
        self.version_history = []
        # Add current user to editors by default
        if hasattr(self, "user_id"):
            self.roles["editor"].append(self.user_id)
        # Version history tracking
        self.save_version()

    def save_version(self):
        """Save current proposal version"""
        self.version_history.append(self.current_proposal.copy())

    def rollback_version(self, version_idx: int):
        """Rollback to a previous version"""
        if 0 <= version_idx < len(self.version_history):
            self.current_proposal = self.version_history[version_idx].copy()
            self.update_preview()

if __name__ == "__main__":
    app = QApplication(sys.argv)
    
    # Test with database
    from ..core.database import setup_database
    setup_database()
    
    editor = ProposalEditorWidget()
    editor.resize(1200, 800)
    editor.show()
    
    sys.exit(app.exec_())
